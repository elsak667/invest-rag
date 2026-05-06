#!/usr/bin/env python3
"""
扫描指定文件夹，解析文档并入库 Supabase documents 表
支持格式：PDF, DOCX, TXT, MD
"""
import subprocess
import base64
import tempfile
import os
import re
from datetime import datetime

# 配置
VPS_URL = "https://nottingham-protected-trivia-cassette.trycloudflare.com/embed"
WORKDIR = "/Users/els/scripts/invest_rag"

def parse_pdf(file_path: str) -> str:
    """解析 PDF 文件"""
    try:
        import pymupdf
        doc = pymupdf.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        print(f"  PDF 解析失败: {e}")
        return ""

def parse_docx(file_path: str) -> str:
    """解析 DOCX 文件（段落 + 表格）"""
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        # 读段落
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        # 读表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(' | '.join(cells))
        return '\n'.join(parts).strip()
    except Exception as e:
        print(f"  DOCX 解析失败: {e}")
        return ""

def parse_txt(file_path: str) -> str:
    """解析纯文本文件"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().strip()
    except Exception as e:
        print(f"  TXT 解析失败: {e}")
        return ""

def extract_company_name(text: str, filename: str) -> str:
    """从内容优先提取公司名，fallback 到文件名"""
    if not text.strip():
        return _name_from_filename(filename)

    first_line = text.strip().split('\n')[0].strip()
    # 去掉 markdown 标题标记
    first_line = re.sub(r'^#+\s*', '', first_line).strip()

    # 1. 如果第一行是 "XXX有限公司/股份" 格式，在正文中搜索真实品牌名
    m_legal = re.search(r'^([^（(]{2,20}?)有限', first_line)
    if m_legal:
        legal_short = m_legal.group(1)  # e.g. "灵明科技"
        prefix = legal_short[:2]  # 取前2字 e.g. "灵明"
        body = text[:1000]
        # 常见品牌后缀（行业词）
        for suffix in ['光子', '光机', '激光', '传感', '芯片', '系统', '机器人']:
            brand = prefix + suffix
            # 从第一个换行后开始找（跳过文件开头可能是法律名的部分）
            first_nl = body.find('\n')
            search_start = first_nl + 1 if first_nl >= 0 else 0
            # 在正文中找 brand 的所有出现
            start = search_start
            while True:
                idx = body.find(brand, start)
                if idx < 0:
                    break
                # 检查前面：若是句首(pos=0)或前字符是标点/空格/换行
                ok_before = (idx == 0) or (body[idx - 1] in ' \t\n\r,，。；;、：:（)）""''【】[]·市省区县镇乡路号栋楼层')
                # 检查后面：若是句尾或后字符是标点/空格/换行/数字，
                # 或是"公司/科技/股份/集团/有限/技术/系统"等公司后缀词
                end_idx = idx + len(brand)
                after_chars = ' \t\n\r,，。；;:、：:）)""''】【]0123456789'
                company_suffixes = ('公司', '科技', '技术', '系统', '股份', '集团', '有限', '有', '为人', '致力于', '成立于', '成立', '设立')
                ok_after = (
                    (end_idx >= len(body)) or
                    (body[end_idx] in after_chars) or
                    any(body[end_idx:].startswith(suf) for suf in company_suffixes)
                )
                if ok_before and ok_after:
                    return brand
                start = idx + 1  # 找下一个出现
        # 备选：legal_short 本身（"灵明科技"）
        return legal_short

    # 2. 去掉括号内标题后缀（如"投资亮点"、"情况报告"）
    clean_line = re.sub(r'[（(].*$', '', first_line).strip()
    industry_kw = any(kw in clean_line for kw in ['光子', '激光', '芯片', '传感', '医疗', '电子', '光机', '量子', '机器人'])
    if 2 <= len(clean_line) <= 18 and industry_kw:
        # 如果 clean_line 是 "行业词 + 空格 + 标题" 格式（如 "灵明光子 投资分析报告"），
        # 只取空格前的第一段（品牌名）
        first_token = clean_line.split()[0] if ' ' in clean_line else clean_line
        if any(kw in first_token for kw in ['光子', '激光', '光机', '传感']):
            return first_token
        # 行业关键词后面不能是另一实词
        for kw in ['光子', '激光', '光机']:
            if clean_line.startswith(kw) and len(clean_line) > len(kw):
                rest = clean_line[len(kw):].strip()
                if not re.match(r'^(科技|技术|系统|股份|有限)', rest):
                    return clean_line
        return clean_line

    # 3. 内容第一行本身就是公司名（如"飞博激光\n"）
    if 2 <= len(first_line) <= 20 and not any(c in first_line for c in '，。：:;；'):
        return first_line

    # 4. 从文件名提取
    return _name_from_filename(filename)


def _name_from_filename(filename: str) -> str:
    """从文件名提取公司名"""
    name = os.path.splitext(os.path.basename(filename))[0]
    for suffix in ['BP', 'bp', '商业计划书', '尽调', '财务', '财报', '报告', 'docx', 'pdf', 'txt', 'md']:
        name = re.sub(rf'{suffix}[_-]?', '', name)
    name = name.strip().replace('_', ' ').replace('-', ' ')
    return name or "未知公司"

def detect_doc_type(text: str, filename: str) -> str:
    """识别文档类型"""
    filename_lower = filename.lower()
    if 'bp' in filename_lower or '商业计划书' in filename_lower:
        return 'bp'
    elif '尽调' in filename_lower or 'due' in filename_lower:
        return 'due_diligence'
    elif '财务' in filename_lower or '财报' in filename_lower:
        return 'financial'
    elif 'cap' in filename_lower or '股权' in filename_lower:
        return 'cap_table'
    elif '合同' in filename_lower or '协议' in filename_lower:
        return 'contract'

    # 从内容判断
    keywords_bp = ['产品', '市场', '团队', '商业模式', '竞争', '壁垒', '融资']
    keywords_financial = ['资产负债表', '利润表', '现金流量', '营收', '净利润', '毛利率']
    keywords_due = ['团队背景', '创始人', '股权结构', '背调', '竞品分析']

    text_lower = text[:2000].lower()
    scores = {
        'bp': sum(1 for k in keywords_bp if k in text_lower),
        'financial': sum(1 for k in keywords_financial if k in text_lower),
        'due_diligence': sum(1 for k in keywords_due if k in text_lower),
    }
    if max(scores.values()) > 0:
        return max(scores, key=scores.get)
    return 'other'

def get_embedding(text: str) -> list[float]:
    """获取文本 embedding"""
    import requests
    try:
        resp = requests.post(VPS_URL, json={"texts": [text]}, timeout=30)
        if resp.status_code == 200:
            return resp.json()["embeddings"][0]
        else:
            print(f"  Embedding 失败: {resp.status_code}")
            return [0.0] * 384
    except Exception as e:
        print(f"  Embedding 请求失败: {e}")
        return [0.0] * 384

def insert_document(company: str, doc_type: str, source: str, content: str, embedding: list[float]) -> bool:
    """写入 documents 表"""
    sql = """
    INSERT INTO documents (company, doc_type, source, content, embedding)
    VALUES ('{company}', '{doc_type}', '{source}', '{content}', '{embedding}');
    """
    # base64 编码避免 SQL 注入
    content_b64 = base64.b64encode(content.encode('utf-8')).decode('ascii')
    emb_str = '[' + ','.join(str(x) for x in embedding) + ']'

    sql = f"""
    INSERT INTO documents (company, doc_type, source, content, embedding)
    VALUES (
        '{company}',
        '{doc_type}',
        '{source}',
        CONVERT_FROM(decode('{content_b64}', 'base64'), 'UTF8'),
        '{emb_str}'::vector
    );
    """

    try:
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w', suffix='.sql', delete=False, encoding='utf-8') as f:
            f.write(sql)
            sql_file = f.name
        r = subprocess.run(
            ['supabase', 'db', 'query', '--linked', '-f', sql_file],
            capture_output=True, text=True, timeout=30, cwd=WORKDIR
        )
        os.unlink(sql_file)
        if r.returncode == 0:
            return True
        else:
            print(f"  写入失败: {r.stderr[:200]}")
            return False
    except Exception as e:
        print(f"  写入异常: {e}")
        return False

def scan_folder(folder_path: str) -> list[str]:
    """扫描文件夹，返回所有支持的文件路径"""
    supported_extensions = ['.pdf', '.docx', '.txt', '.md']
    files = []
    for root, dirs, filenames in os.walk(folder_path):
        for filename in filenames:
            ext = os.path.splitext(filename)[1].lower()
            if ext in supported_extensions:
                files.append(os.path.join(root, filename))
    return files

def process_file(file_path: str) -> dict | None:
    """处理单个文件，返回解析结果"""
    print(f"\n处理: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        text = parse_pdf(file_path)
    elif ext == '.docx':
        text = parse_docx(file_path)
    else:
        text = parse_txt(file_path)

    if not text or len(text) < 100:
        print(f"  ⚠️  内容过短，跳过")
        return None

    # 截取前 5000 字用于 embedding（避免 token 溢出）
    text_short = text[:5000]

    company = extract_company_name(text, os.path.basename(file_path))
    doc_type = detect_doc_type(text, os.path.basename(file_path))

    print(f"  公司: {company}")
    print(f"  类型: {doc_type}")
    print(f"  字数: {len(text)}")

    print(f"  生成 embedding...")
    embedding = get_embedding(text_short)

    print(f"  写入数据库...")
    ok = insert_document(company, doc_type, file_path, text, embedding)

    if ok:
        print(f"  ✅ 成功")
        return {"company": company, "doc_type": doc_type, "file": file_path, "chars": len(text)}
    else:
        print(f"  ❌ 失败")
        return None

def main():
    import sys

    if len(sys.argv) < 2:
        print("用法: python scan_documents.py <文件夹路径>")
        print("示例: python scan_documents.py /Users/els/Desktop/investment_docs")
        sys.exit(1)

    folder_path = sys.argv[1]

    if not os.path.isdir(folder_path):
        print(f"❌ 目录不存在: {folder_path}")
        sys.exit(1)

    files = scan_folder(folder_path)
    if not files:
        print(f"❌ 未找到支持的文档 (PDF/DOCX/TXT/MD)")
        sys.exit(1)

    print(f"\n{'='*60}")
    print(f"文档扫描")
    print(f"{'='*60}")
    print(f"目录: {folder_path}")
    print(f"找到文件: {len(files)} 个\n")

    results = []
    for i, file_path in enumerate(files, 1):
        print(f"\n[{i}/{len(files)}]")
        result = process_file(file_path)
        if result:
            results.append(result)

    print(f"\n{'='*60}")
    print(f"扫描完成: {len(results)}/{len(files)} 个文件入库")
    print(f"{'='*60}")

    if results:
        print("\n入库文件列表:")
        for r in results:
            print(f"  [{r['doc_type']}] {r['company']} - {os.path.basename(r['file'])}")

if __name__ == '__main__':
    main()
